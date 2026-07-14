from __future__ import annotations

from pathlib import Path
import html
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "feishu-share-script.md"
OUTPUT = ROOT / "feishu-share-script.docx"
OUTPUT_HTML = ROOT / "feishu-share-script.html"


def set_run_font(run, name: str = "Arial", size: int | None = None, bold: bool | None = None, color: tuple[int, int, int] | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)

    heading_specs = {
        "Heading 1": (16, 18, 8),
        "Heading 2": (14, 14, 6),
        "Heading 3": (12, 10, 4),
    }
    for style_name, (size, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=11)


def add_metadata(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11, color=(85, 85, 85))


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=11)


def add_image(doc: Document, relative_path: str, alt_text: str) -> None:
    image_path = ROOT / relative_path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    set_run_font(run)
    caption = doc.add_paragraph(style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)
    cap_run = caption.add_run(alt_text)
    set_run_font(cap_run, size=10, color=(85, 85, 85))


def build() -> None:
    doc = Document()
    configure_styles(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    image_pattern = re.compile(r"!\[(?P<alt>.*?)]\((?P<path>.*?)\)")

    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line[2:].strip())
            set_run_font(run, size=24, bold=True)
            continue
        if line.startswith("分享人：") or line.startswith("时间："):
            add_metadata(doc, line)
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line[3:].strip())
            set_run_font(run, size=16, bold=True)
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line[4:].strip())
            set_run_font(run, size=14, bold=True)
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            continue
        match = image_pattern.fullmatch(line)
        if match:
            add_image(doc, match.group("path"), match.group("alt"))
            continue
        add_paragraph(doc, line.replace("  ", " "))

    doc.save(OUTPUT)


def build_html() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    image_pattern = re.compile(r"!\[(?P<alt>.*?)]\((?P<path>.*?)\)")
    parts: list[str] = [
        "<!doctype html>",
        '<html lang="zh-CN">',
        "<head>",
        '<meta charset="utf-8">',
        "<title>Multica + Trellis：AI Agent 协作开发实践分享</title>",
        "<style>",
        "body { margin: 0; background: #f5f7fb; font-family: Arial, 'PingFang SC', 'Hiragino Sans GB', 'Microsoft YaHei', sans-serif; color: #111; }",
        ".page { width: 900px; margin: 32px auto; background: #fff; padding: 56px 72px 72px; box-shadow: 0 10px 30px rgba(15, 23, 42, 0.08); border-radius: 18px; }",
        "h1 { font-size: 36px; line-height: 1.25; margin: 0 0 12px; }",
        "h2 { font-size: 26px; line-height: 1.35; margin: 34px 0 14px; }",
        "h3 { font-size: 20px; line-height: 1.4; margin: 24px 0 10px; }",
        "p { font-size: 17px; line-height: 1.8; margin: 0 0 14px; }",
        ".meta { font-size: 15px; color: #555; margin-bottom: 4px; }",
        "figure { margin: 22px 0 24px; }",
        "figure img { width: 100%; border-radius: 12px; border: 1px solid #e5e7eb; }",
        "figcaption { font-size: 14px; color: #666; text-align: center; margin-top: 8px; }",
        "ul { margin: 0 0 16px 24px; padding: 0; }",
        "li { font-size: 17px; line-height: 1.8; margin: 0 0 6px; }",
        "</style>",
        "</head>",
        "<body>",
        '<main class="page">',
    ]

    in_list = False
    for raw in lines:
        line = raw.rstrip()
        if not line:
            if in_list:
                parts.append("</ul>")
                in_list = False
            continue
        match = image_pattern.fullmatch(line)
        if match:
            if in_list:
                parts.append("</ul>")
                in_list = False
            alt = html.escape(match.group("alt"))
            path = html.escape(match.group("path"))
            parts.append(f'<figure><img src="{path}" alt="{alt}"><figcaption>{alt}</figcaption></figure>')
            continue
        if line.startswith("# "):
            if in_list:
                parts.append("</ul>")
                in_list = False
            parts.append(f"<h1>{html.escape(line[2:].strip())}</h1>")
            continue
        if line.startswith("分享人：") or line.startswith("时间："):
            if in_list:
                parts.append("</ul>")
                in_list = False
            parts.append(f'<p class="meta">{html.escape(line)}</p>')
            continue
        if line.startswith("## "):
            if in_list:
                parts.append("</ul>")
                in_list = False
            parts.append(f"<h2>{html.escape(line[3:].strip())}</h2>")
            continue
        if line.startswith("### "):
            if in_list:
                parts.append("</ul>")
                in_list = False
            parts.append(f"<h3>{html.escape(line[4:].strip())}</h3>")
            continue
        if line.startswith("- "):
            if not in_list:
                parts.append("<ul>")
                in_list = True
            parts.append(f"<li>{html.escape(line[2:].strip())}</li>")
            continue
        if in_list:
            parts.append("</ul>")
            in_list = False
        parts.append(f"<p>{html.escape(line)}</p>")

    if in_list:
        parts.append("</ul>")
    parts.extend(["</main>", "</body>", "</html>"])
    OUTPUT_HTML.write_text("\n".join(parts), encoding="utf-8")


if __name__ == "__main__":
    build()
    build_html()
